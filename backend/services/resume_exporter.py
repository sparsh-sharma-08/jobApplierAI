"""
Resume export to DOCX and PDF formats.
"""
import json
import os
from typing import Dict, Any, Optional
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./data/resumes")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def _clean_list(lst):
    if not isinstance(lst, list): return []
    cleaned = []
    for item in lst:
        if isinstance(item, dict):
            vals = list(item.values())
            cleaned.append(str(vals[0]) if vals else "")
        elif item is not None:
            cleaned.append(str(item))
    return cleaned


def clean_resume_data(data: Dict) -> Dict:
    def _clean_str(val):
        return str(val).strip() if val is not None else ""

    data["name"] = _clean_str(data.get("name"))
    data["summary"] = _clean_str(data.get("summary"))
    
    if "contact" in data and isinstance(data["contact"], dict):
        for k, v in data["contact"].items():
            data["contact"][k] = _clean_str(v)
            
    data["skills"] = _clean_list(data.get("skills", []))
    data["certifications"] = _clean_list(data.get("certifications", []))
    data["languages"] = _clean_list(data.get("languages", []))
    
    for exp in data.get("experience", []):
        if isinstance(exp, dict):
            for k in ["title", "company", "start_date", "end_date"]:
                exp[k] = _clean_str(exp.get(k))
            exp["highlights"] = _clean_list(exp.get("highlights", []))
            exp["technologies"] = _clean_list(exp.get("technologies", []))
            
    for proj in data.get("projects", []):
        if isinstance(proj, dict):
            for k in ["name", "description", "link"]:
                proj[k] = _clean_str(proj.get(k))
            proj["highlights"] = _clean_list(proj.get("highlights", []))
            proj["technologies"] = _clean_list(proj.get("technologies", []))
            
    for edu in data.get("education", []):
        if isinstance(edu, dict):
            for k in ["degree", "institution", "year", "gpa"]:
                edu[k] = _clean_str(edu.get(k))

    return data


def save_resume_json(resume_data: Dict, job_id: int) -> str:
    path = os.path.join(OUTPUT_DIR, f"resume_job_{job_id}_{_ts()}.json")
    with open(path, "w") as f:
        json.dump(resume_data, f, indent=2)
    return path


def save_resume_docx(resume_data: Dict, job_id: int) -> Optional[str]:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Narrow margins for ATS
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(resume_data.get("name", ""))
        run.bold = True
        run.font.size = Pt(18)

        # Contact
        contact = resume_data.get("contact", {})
        contact_items = [v for v in [contact.get("email"), contact.get("phone"),
                                      contact.get("location"), contact.get("linkedin"),
                                      contact.get("github")] if v]
        if contact_items:
            contact_para = doc.add_paragraph(" | ".join(contact_items))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if contact_para.runs:
                contact_para.runs[0].font.size = Pt(9)

        def add_section_heading(title):
            p = doc.add_paragraph()
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(10)
            # Add horizontal rule via bottom border
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)

        # Summary
        summary = resume_data.get("summary")
        if summary:
            add_section_heading("Professional Summary")
            doc.add_paragraph(summary).runs[0].font.size = Pt(10)

        # Skills
        skills = resume_data.get("skills", [])
        if skills:
            add_section_heading("Technical Skills")
            doc.add_paragraph(", ".join(skills)).runs[0].font.size = Pt(10)

        # Experience
        experience = resume_data.get("experience", [])
        if experience:
            add_section_heading("Experience")
            for exp in experience:
                p = doc.add_paragraph()
                r = p.add_run(f"{exp.get('title', '')} — {exp.get('company', '')}")
                r.bold = True
                r.font.size = Pt(10)
                dates = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
                p2 = doc.add_paragraph(dates)
                p2.runs[0].italic = True
                p2.runs[0].font.size = Pt(9)
                for highlight in exp.get("highlights", []):
                    doc.add_paragraph(f"• {highlight}", style="List Bullet").runs[0].font.size = Pt(10)

        # Projects
        projects = resume_data.get("projects", [])
        if projects:
            add_section_heading("Projects")
            for proj in projects:
                p = doc.add_paragraph()
                r = p.add_run(proj.get("name", ""))
                r.bold = True
                r.font.size = Pt(10)
                tech = proj.get("technologies", [])
                if tech:
                    p.add_run(f"  ({', '.join(tech)})").font.size = Pt(9)
                desc = proj.get("description", "")
                if desc:
                    desc_para = doc.add_paragraph(desc)
                    if desc_para.runs:
                        desc_para.runs[0].font.size = Pt(10)
                for h in proj.get("highlights", []):
                    h_para = doc.add_paragraph(f"• {h}", style="List Bullet")
                    if h_para.runs:
                        h_para.runs[0].font.size = Pt(10)

        # Education
        education = resume_data.get("education", [])
        if education:
            add_section_heading("Education")
            for edu in education:
                p = doc.add_paragraph()
                r = p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}")
                r.bold = True
                r.font.size = Pt(10)
                details = edu.get("year", "")
                if edu.get("gpa"):
                    details += f"  |  GPA: {edu['gpa']}"
                if details:
                    doc.add_paragraph(details).runs[0].font.size = Pt(9)

        # Certifications
        certs = resume_data.get("certifications", [])
        if certs:
            add_section_heading("Certifications")
            for cert in certs:
                doc.add_paragraph(f"• {cert}").runs[0].font.size = Pt(10)

        path = os.path.join(OUTPUT_DIR, f"resume_job_{job_id}_{_ts()}.docx")
        doc.save(path)
        return path

    except ImportError:
        logger.warning("python-docx not installed. Skipping DOCX export.")
        return None
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        return None


def save_resume_pdf(docx_path: Optional[str], resume_data: Dict, job_id: int) -> Optional[str]:
    """Generate a clean, ATS-compatible PDF matching the reference resume format."""
    pdf_path = os.path.join(OUTPUT_DIR, f"resume_job_{job_id}_{_ts()}.pdf")

    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

        doc_pdf = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
            leftMargin=0.6 * inch,
            rightMargin=0.6 * inch
        )

        # ── Styles ──────────────────────────────────────────────
        name_style = ParagraphStyle(
            'Name', fontName='Helvetica-Bold', fontSize=20, alignment=TA_CENTER,
            spaceAfter=2, leading=24
        )
        contact_style = ParagraphStyle(
            'Contact', fontName='Helvetica', fontSize=9, alignment=TA_CENTER,
            spaceAfter=4, leading=12
        )
        section_heading_style = ParagraphStyle(
            'SectionHeading', fontName='Helvetica-Bold', fontSize=11, 
            spaceBefore=6, spaceAfter=2, leading=14
        )
        body_style = ParagraphStyle(
            'Body', fontName='Helvetica', fontSize=10, alignment=TA_JUSTIFY,
            spaceAfter=4, leading=13
        )
        project_title_style = ParagraphStyle(
            'ProjectTitle', fontName='Helvetica-Bold', fontSize=10.5,
            spaceAfter=1, spaceBefore=4, leading=13
        )
        bullet_style = ParagraphStyle(
            'Bullet', fontName='Helvetica', fontSize=10, leftIndent=18,
            bulletIndent=8, spaceAfter=2, leading=13, alignment=TA_JUSTIFY
        )
        tech_style = ParagraphStyle(
            'Tech', fontName='Helvetica', fontSize=10, leftIndent=18,
            spaceAfter=4, leading=12, textColor=colors.HexColor('#222222')
        )
        edu_left_style = ParagraphStyle(
            'EduLeft', fontName='Helvetica-Bold', fontSize=10.5, leading=14
        )
        edu_right_style = ParagraphStyle(
            'EduRight', fontName='Helvetica-Bold', fontSize=10.5, alignment=TA_RIGHT, leading=14
        )
        edu_detail_style = ParagraphStyle(
            'EduDetail', fontName='Helvetica', fontSize=10, leftIndent=18, spaceAfter=2, leading=12
        )

        story = []

        # ── NAME ────────────────────────────────────────────────
        story.append(Paragraph(resume_data.get("name", "").upper(), name_style))

        # ── CONTACT LINE ────────────────────────────────────────
        contact = resume_data.get("contact", {})
        contact_parts = []
        if contact.get("email"): contact_parts.append(contact["email"])
        if contact.get("phone"): contact_parts.append(contact["phone"])
        if contact.get("location"): contact_parts.append(contact["location"])
        
        # Add links if available
        links = contact.get("links", [])
        if links:
            # just take the first two to keep it clean, usually LinkedIn and GitHub
            contact_parts.extend(links[:2])

        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), contact_style))

        # ── SUMMARY ─────────────────────────────────────────────
        summary = resume_data.get("summary", "")
        if summary:
            story.append(Paragraph("<b>SUMMARY</b>", section_heading_style))
            add_hr()
            story.append(Paragraph(summary, body_style))

        # ── SKILLS ──────────────────────────────────────────────
        skills = resume_data.get("skills", [])
        if skills:
            story.append(Paragraph("<b>SKILLS</b>", section_heading_style))
            add_hr()
            # join all skills with a • character in a single paragraph for compactness
            skills_str = " • ".join(skills)
            story.append(Paragraph(skills_str, body_style))

        # ── EXPERIENCE ──────────────────────────────────────────
        experience = resume_data.get("experience", [])
        if experience:
            story.append(Paragraph("<b>PROFESSIONAL EXPERIENCE</b>", section_heading_style))
            for exp in experience:
                title = exp.get("title", "")
                company = exp.get("company", "")
                date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}".strip(" -")
                location = exp.get("location", "")
                highlights = exp.get("highlights", [])
                
                # Title and Date on same line using table
                header_data = [
                    [Paragraph(f"<b>{title}</b>", ParagraphStyle('T', fontName='Helvetica-Bold', fontSize=10.5)), 
                     Paragraph(date_range, ParagraphStyle('D', fontName='Helvetica', fontSize=10, alignment=TA_RIGHT))]
                ]
                t = Table(header_data, colWidths=[4*inch, 3.3*inch])
                t.setStyle(TableStyle([
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 2),
                ]))
                story.append(t)

                # Company and Location
                comp_loc = company
                if location: comp_loc += f", {location}"
                story.append(Paragraph(f"<i>{comp_loc}</i>", ParagraphStyle('C', fontName='Helvetica-Oblique', fontSize=10, spaceAfter=3)))

                # Highlights (Bullets)
                for h in highlights:
                    # Bold the first word (action verb)
                    parts = h.split(" ", 1)
                    if len(parts) == 2:
                        bullet_text = f"<b>{parts[0]}</b> {parts[1]}"
                    else:
                        bullet_text = h
                    story.append(Paragraph(f"• {bullet_text}", bullet_style))
                
                techs = exp.get("technologies", [])
                if techs:
                    story.append(Paragraph(f"<b>Tech:</b> {', '.join(techs)}", tech_style))

            add_hr()

        # ── PROJECTS ────────────────────────────────────────────
        projects = resume_data.get("projects", [])
        if projects:
            story.append(Paragraph("<b>PROJECTS</b>", section_heading_style))
            for proj in projects:
                name = proj.get("name", "")
                desc = proj.get("description", "")
                techs = proj.get("technologies", [])
                highlights = proj.get("highlights", [])

                # Project Name (bold)
                story.append(Paragraph(f"<b>{name}</b>", project_title_style))
                
                if desc:
                    story.append(Paragraph(desc, bullet_style))
                
                for h in highlights:
                    # Bold the first word (action verb)
                    parts = h.split(" ", 1)
                    if len(parts) == 2:
                        bullet_text = f"<b>{parts[0]}</b> {parts[1]}"
                    else:
                        bullet_text = h
                    story.append(Paragraph(f"• {bullet_text}", bullet_style))

                if techs:
                    story.append(Paragraph(f"<b>Tech:</b> {', '.join(techs)}", tech_style))

            add_hr()

        # ── EDUCATION ───────────────────────────────────────────
        education = resume_data.get("education", [])
        if education:
            story.append(Paragraph("<b>EDUCATION</b>", section_heading_style))
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")
                gpa = edu.get("gpa", "")
                
                # Degree + Institution on left, Year on right  
                left_text = f"<b>{degree}</b>"
                if institution:
                    left_text += f" — {institution}"
                
                edu_table = Table(
                    [[Paragraph(left_text, ParagraphStyle('EduLeft', fontName='Helvetica-Bold', fontSize=10.5, leading=14)), 
                      Paragraph(f"<b>{year}</b>", ParagraphStyle('EduRight', fontName='Helvetica-Bold', fontSize=10.5, alignment=TA_RIGHT, leading=14))]],
                    colWidths=[(doc_pdf.width * 0.75), (doc_pdf.width * 0.25)]
                )
                edu_table.setStyle(TableStyle([
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ]))
                story.append(edu_table)
                
                if gpa:
                    story.append(Paragraph(f"• GPA: {gpa}", ParagraphStyle('EduDetail', fontName='Helvetica', fontSize=10, leftIndent=18, spaceAfter=2, leading=12)))
            
            add_hr()

        # ── CERTIFICATIONS / ACHIEVEMENTS ───────────────────────
        certs = resume_data.get("certifications", [])
        if certs:
            story.append(Paragraph("<b>ACHIEVEMENTS</b>", section_heading_style))
            for cert in certs:
                story.append(Paragraph(f"• {cert}", bullet_style))

        doc_pdf.build(story)
        pdf_bytes = pdf_buffer.getvalue()
        pdf_buffer.close()
        return pdf_bytes

    except ImportError:
        logger.warning("reportlab not installed. Skipping PDF export.")
        return None
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return None


def _ts():
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def export_resume(resume_data: Dict, job_id: int) -> Dict[str, Optional[str]]:
    """Export resume to all formats. Returns dict of paths."""
    resume_data = clean_resume_data(resume_data)
    json_path = save_resume_json(resume_data, job_id)
    docx_path = save_resume_docx(resume_data, job_id)
    pdf_path = save_resume_pdf(docx_path, resume_data, job_id)
    return {
        "json": json_path,
        "docx": docx_path,
        "pdf": pdf_path
    }

def generate_resume_pdf_bytes(resume_data: Dict) -> Optional[bytes]:
    """Generate a clean, ATS-compatible PDF in memory matching the reference resume format."""
    import io
    pdf_buffer = io.BytesIO()

    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

        doc_pdf = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
            leftMargin=0.6 * inch,
            rightMargin=0.6 * inch
        )

        name_style = ParagraphStyle(
            'Name', fontName='Helvetica-Bold', fontSize=20, alignment=TA_CENTER,
            spaceAfter=2, leading=24
        )
        contact_style = ParagraphStyle(
            'Contact', fontName='Helvetica', fontSize=9, alignment=TA_CENTER,
            spaceAfter=4, leading=12
        )
        section_heading_style = ParagraphStyle(
            'SectionHeading', fontName='Helvetica-Bold', fontSize=11, 
            spaceBefore=6, spaceAfter=2, leading=14
        )
        body_style = ParagraphStyle(
            'Body', fontName='Helvetica', fontSize=10, alignment=TA_JUSTIFY,
            spaceAfter=4, leading=13
        )
        project_title_style = ParagraphStyle(
            'ProjectTitle', fontName='Helvetica-Bold', fontSize=10.5,
            spaceAfter=1, spaceBefore=4, leading=13
        )
        bullet_style = ParagraphStyle(
            'Bullet', fontName='Helvetica', fontSize=10, leftIndent=18,
            firstLineIndent=-10, spaceAfter=2, leading=13
        )
        tech_style = ParagraphStyle(
            'Tech', fontName='Helvetica-Oblique', fontSize=9, leftIndent=18,
            spaceBefore=1, spaceAfter=4, textColor=colors.grey
        )

        story = []

        def add_hr():
            story.append(Spacer(1, 2))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceBefore=1, spaceAfter=4))

        # Header
        name = resume_data.get("name", "Candidate Name")
        story.append(Paragraph(f"<b>{name}</b>", name_style))
        
        contact = resume_data.get("contact", {})
        contact_parts = []
        if contact.get("email"): contact_parts.append(contact["email"])
        if contact.get("phone"): contact_parts.append(contact["phone"])
        if contact.get("location"): contact_parts.append(contact["location"])
        
        links = contact.get("links", [])
        if links:
            contact_parts.extend(links[:2])

        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), contact_style))

        # Summary
        summary = resume_data.get("summary", "")
        if summary:
            story.append(Paragraph("<b>SUMMARY</b>", section_heading_style))
            add_hr()
            story.append(Paragraph(summary, body_style))

        # Skills
        skills = resume_data.get("skills", [])
        if skills:
            story.append(Paragraph("<b>SKILLS</b>", section_heading_style))
            add_hr()
            skills_str = " • ".join(skills)
            story.append(Paragraph(skills_str, body_style))

        # Experience
        experience = resume_data.get("experience", [])
        if experience:
            story.append(Paragraph("<b>PROFESSIONAL EXPERIENCE</b>", section_heading_style))
            for exp in experience:
                title = exp.get("title", "")
                company = exp.get("company", "")
                date_range = f"{exp.get('start_date', '')} - {exp.get('end_date', 'Present')}".strip(" -")
                location = exp.get("location", "")
                highlights = exp.get("highlights", [])
                
                header_data = [
                    [Paragraph(f"<b>{title}</b>", ParagraphStyle('T', fontName='Helvetica-Bold', fontSize=10.5)), 
                     Paragraph(date_range, ParagraphStyle('D', fontName='Helvetica', fontSize=10, alignment=TA_RIGHT))]
                ]
                t = Table(header_data, colWidths=[4*inch, 3.3*inch])
                t.setStyle(TableStyle([
                    ('LEFTPADDING', (0,0), (-1,-1), 0),
                    ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 2),
                ]))
                story.append(t)

                comp_loc = company
                if location: comp_loc += f", {location}"
                story.append(Paragraph(f"<i>{comp_loc}</i>", ParagraphStyle('C', fontName='Helvetica-Oblique', fontSize=10, spaceAfter=3)))

                for h in highlights:
                    parts = h.split(" ", 1)
                    bullet_text = f"<b>{parts[0]}</b> {parts[1]}" if len(parts) == 2 else h
                    story.append(Paragraph(f"• {bullet_text}", bullet_style))
                
                techs = exp.get("technologies", [])
                if techs:
                    story.append(Paragraph(f"<b>Tech:</b> {', '.join(techs)}", tech_style))

            add_hr()

        # Projects
        projects = resume_data.get("projects", [])
        if projects:
            story.append(Paragraph("<b>PROJECTS</b>", section_heading_style))
            for proj in projects:
                name = proj.get("name", "")
                desc = proj.get("description", "")
                techs = proj.get("technologies", [])
                highlights = proj.get("highlights", [])

                story.append(Paragraph(f"<b>{name}</b>", project_title_style))
                if desc: story.append(Paragraph(desc, bullet_style))
                
                for h in highlights:
                    parts = h.split(" ", 1)
                    bullet_text = f"<b>{parts[0]}</b> {parts[1]}" if len(parts) == 2 else h
                    story.append(Paragraph(f"• {bullet_text}", bullet_style))

                if techs:
                    story.append(Paragraph(f"<b>Tech:</b> {', '.join(techs)}", tech_style))

            add_hr()

        # Education
        education = resume_data.get("education", [])
        if education:
            story.append(Paragraph("<b>EDUCATION</b>", section_heading_style))
            for edu in education:
                degree = edu.get("degree", "")
                institution = edu.get("institution", "")
                year = edu.get("year", "")
                gpa = edu.get("gpa", "")

                left_text = f"<b>{degree}</b>"
                if institution:
                    left_text += f" — {institution}"
                
                edu_table = Table(
                    [[Paragraph(left_text, ParagraphStyle('EduL', fontName='Helvetica-Bold', fontSize=10.5, leading=14)), 
                      Paragraph(f"<b>{year}</b>", ParagraphStyle('EduR', fontName='Helvetica-Bold', fontSize=10.5, alignment=TA_RIGHT, leading=14))]],
                    colWidths=[4*inch, 3.3*inch]
                )
                edu_table.setStyle(TableStyle([
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0),
                    ('TOPPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 1),
                ]))
                story.append(edu_table)
                if gpa:
                    story.append(Paragraph(f"• GPA: {gpa}", ParagraphStyle('GPA', fontName='Helvetica', fontSize=10, leftIndent=18, spaceAfter=2, leading=12)))
            
            add_hr()

        # Certifications
        certs = resume_data.get("certifications", [])
        if certs:
            story.append(Paragraph("<b>ACHIEVEMENTS</b>", section_heading_style))
            for cert in certs:
                story.append(Paragraph(f"• {cert}", bullet_style))

        doc_pdf.build(story)
        pdf_bytes = pdf_buffer.getvalue()
        pdf_buffer.close()
        return pdf_bytes

    except ImportError:
        logger.warning("reportlab not installed.")
        return None
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return None

